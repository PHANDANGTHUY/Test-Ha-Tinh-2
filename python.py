import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
import re
import numpy_financial as npf
import google.generativeai as genai
from io import BytesIO
from datetime import datetime
# ======================================================================================
# CẤU HÌNH TRANG VÀ KHỞI TẠO
# ======================================================================================
st.set_page_config(
    page_title="Hệ thống Thẩm định Phương án Kinh doanh",
    page_icon="🏦",
    layout="wide",
)
st.title("🏦 Hệ thống Thẩm định Phương án Kinh doanh")
st.markdown("---")
# Khởi tạo session state để lưu trữ dữ liệu giữa các lần re-run
if 'data_extracted' not in st.session_state:
    st.session_state.data_extracted = False
if 'docx_text' not in st.session_state:
    st.session_state.docx_text = ""
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
# ======================================================================================
# CÁC HÀM HỖ TRỢ (HELPERS)
# ======================================================================================
def format_currency(value, decimal_places=0):
    """Định dạng số thành chuỗi tiền tệ với dấu chấm phân cách hàng nghìn."""
    if value is None or not isinstance(value, (int, float)):
        return "0"
    return f"{value:,.{decimal_places}f}".replace(",", ".")
def extract_text_from_docx(docx_file):
    """Trích xuất toàn bộ văn bản từ file .docx."""
    try:
        doc = Document(docx_file)
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Lỗi khi đọc file .docx: {e}")
        return ""
def parse_info_from_text(text):
    """
    Phân tích văn bản để trích xuất thông tin ban đầu (best-effort).
    Hàm này sử dụng regex đơn giản và giả định cấu trúc file.
    """
    info = {}
   
    # Hàm tìm kiếm an toàn
    def safe_search(pattern, text, group=1, default=None, is_numeric=False):
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            result = match.group(group).strip()
            if is_numeric:
                try:
                    # Loại bỏ ký tự không phải số (giữ lại dấu phẩy/chấm)
                    cleaned_result = re.sub(r'[^\d,.]', '', result)
                    # Chuẩn hóa về định dạng số của Python
                    cleaned_result = cleaned_result.replace('.', '').replace(',', '.')
                    return float(cleaned_result)
                except (ValueError, IndexError):
                    return default
            return result
        return default
    # Thông tin khách hàng (không có trong file, giữ mặc định)
    info['ho_ten'] = safe_search(r"Họ và tên:\s*(.*?)\s*\.", text) or "Chưa rõ"
    info['cccd'] = safe_search(r"CCCD số:\s*(\d+)", text) or ""
    info['dia_chi'] = safe_search(r"Nơi cư trú:\s*(.*?)(?:,|$|\n)", text) or "Chưa rõ"
    info['sdt'] = safe_search(r"Số điện thoại:\s*(\d+)", text) or ""
    # Thông tin phương án vay
    info['muc_dich_vay'] = safe_search(r"Mục đích vay:\s*(.*?)\n", text) or "Kinh doanh vật liệu xây dựng"
    info['tong_nhu_cau_von'] = safe_search(r"- Chi phí kinh doanh:\s*([\d.,]+)\s*đồng", text, is_numeric=True, default=0)
    info['von_doi_ung'] = safe_search(r"Vốn đối ứng.*?đồng,([\d.,]+)", text, is_numeric=True, default=0)  # Không có trong file, giữ default
    info['so_tien_vay'] = safe_search(r"Chênh lệch thu chi:\s*([\d.,]+)\s*đồng", text, is_numeric=True, default=0)  # Sử dụng chênh lệch làm proxy nếu cần
    info['lai_suat'] = safe_search(r"Lãi suất đề nghị:\s*([\d.,]+)%/năm", text, is_numeric=True, default=5.0)
    info['thoi_gian_vay'] = safe_search(r"Thời hạn cho vay:\s*(\d+)\s*tháng", text, is_numeric=True, default=3)
    # Thông tin tài sản đảm bảo (không có chi tiết cụ thể)
    info['tsdb_mo_ta'] = safe_search(r"Tài sản bảo đảm:\s*(.*?)(?=III\. Thông tin)", text) or "Chưa có mô tả"
    info['tsdb_gia_tri'] = safe_search(r"Tổng tài sản đảm bảo:\s*([\d.,]+)", text, is_numeric=True, default=0)
    # Trích xuất thêm thông tin cụ thể từ file
    info['doanh_thu'] = safe_search(r"\+Doanh thu của phương án:\s*([\d.,]+)\s*đồng", text, is_numeric=True, default=0)
    info['chi_phi'] = safe_search(r"\+  Chi phí kinh doanh:\s*([\d.,]+)\s*đồng", text, is_numeric=True, default=0)
    info['chenh_lech_thu_chi'] = safe_search(r"\+  Chênh lệch thu chi:\s*([\d.,]+)\s*đồng", text, is_numeric=True, default=0)
    info['nguon_tra_no'] = safe_search(r"- Từ nguồn thu của phương án kinh doanh:\s*([\d.,]+)đồng", text, is_numeric=True, default=0)
    return info
def calculate_repayment_schedule(principal, annual_rate, term_months):
    """Tạo bảng kế hoạch trả nợ chi tiết."""
    if not all([principal > 0, annual_rate > 0, term_months > 0]):
        return pd.DataFrame()
    monthly_rate = annual_rate / 12 / 100
   
    # Tính tiền gốc phải trả hàng tháng
    principal_payment = principal / term_months
   
    schedule = []
    remaining_balance = principal
    for i in range(1, term_months + 1):
        interest_payment = remaining_balance * monthly_rate
        total_payment = principal_payment + interest_payment
       
        schedule.append({
            "Kỳ trả nợ": i,
            "Dư nợ đầu kỳ": remaining_balance,
            "Gốc trả trong kỳ": principal_payment,
            "Lãi trả trong kỳ": interest_payment,
            "Tổng gốc và lãi": total_payment,
            "Dư nợ cuối kỳ": remaining_balance - principal_payment,
        })
        remaining_balance -= principal_payment
    df = pd.DataFrame(schedule)
    # Đảm bảo dư nợ cuối kỳ cuối cùng là 0
    df.loc[df.index[-1], 'Dư nợ cuối kỳ'] = 0
    return df
def generate_excel_download(df):
    """Tạo file Excel từ DataFrame để tải xuống."""
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='LichTraNo')
    processed_data = output.getvalue()
    return processed_data
def generate_report_docx(customer_info, loan_info, collateral_info, ratios, ai_analysis_1, ai_analysis_2):
    """Tạo file Báo cáo Thẩm định dưới dạng .docx."""
    doc = Document()
    doc.add_heading('BÁO CÁO THẨM ĐỊNH SƠ BỘ', level=1)
    doc.add_paragraph(f"Ngày lập báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    # Thông tin khách hàng
    doc.add_heading('1. Thông tin khách hàng', level=2)
    for key, value in customer_info.items():
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
   
    # Thông tin khoản vay
    doc.add_heading('2. Thông tin Phương án vay & Các chỉ số', level=2)
    for key, value in loan_info.items():
        val_str = format_currency(value) if isinstance(value, (int, float)) and 'lãi suất' not in key and 'thời gian' not in key else value
        unit = " %/năm" if 'lãi suất' in key else " tháng" if 'thời gian' in key else " VNĐ" if isinstance(value, (int, float)) else ""
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {val_str}{unit}")
    for key, value in ratios.items():
        doc.add_paragraph(f"{key}: {value}")
   
    # Tài sản đảm bảo
    doc.add_heading('3. Thông tin tài sản đảm bảo', level=2)
    doc.add_paragraph(f"Mô tả: {collateral_info['tsdb_mo_ta']}")
    doc.add_paragraph(f"Tổng giá trị định giá: {format_currency(collateral_info['tsdb_gia_tri'])} VNĐ")
    # Phân tích từ AI
    doc.add_heading('4. Phân tích tự động bởi AI', level=2)
    doc.add_heading('4.1. Phân tích từ file .docx của khách hàng', level=3)
    doc.add_paragraph(ai_analysis_1 if ai_analysis_1 else "Chưa có phân tích.")
    doc.add_heading('4.2. Phân tích từ dữ liệu đã hiệu chỉnh trên ứng dụng', level=3)
    doc.add_paragraph(ai_analysis_2 if ai_analysis_2 else "Chưa có phân tích.")
    # Lưu vào buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
# ======================================================================================
# THANH BÊN (SIDEBAR)
# ======================================================================================
with st.sidebar:
    st.header("Cấu hình & Chức năng")
   
    # Nhập API Key
    api_key = st.text_input("🔑 Gemini API Key", type="password", help="Nhập API Key của bạn để kích hoạt các tính năng AI.")
   
    st.markdown("---")
   
    # Chức năng xuất dữ liệu
    st.subheader("Chức năng Xuất dữ liệu")
    export_option = st.selectbox(
        "Chọn loại báo cáo:",
        ("--- Chọn ---", "Xuất Kế hoạch trả nợ (Excel)", "Xuất Báo cáo Thẩm định")
    )
    execute_export = st.button("Thực hiện", use_container_width=True, disabled=(export_option == "--- Chọn ---"))
# ======================================================================================
# KHU VỰC CHỨC NĂNG CHÍNH (TABS)
# ======================================================================================
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📝 Nhập liệu & Trích xuất thông tin",
    "📊 Phân tích Chỉ số & Dòng tiền",
    "📈 Biểu đồ Trực quan",
    "🤖 Phân tích bởi AI",
    "💬 Chatbot Hỗ trợ"
])
# --------------------------------------------------------------------------------------
# TAB 1: NHẬP LIỆU & TRÍCH XUẤT
# --------------------------------------------------------------------------------------
with tab1:
    st.header("Tải lên và Hiệu chỉnh Thông tin Phương án Kinh doanh")
   
    uploaded_file = st.file_uploader(
        "Tải lên file Phương án Kinh doanh của khách hàng (.docx)",
        type=['docx'],
        accept_multiple_files=False
    )
    if uploaded_file is not None:
        # Chỉ trích xuất lại nếu file thay đổi hoặc chưa trích xuất
        if not st.session_state.data_extracted:
            with st.spinner("Đang đọc và trích xuất thông tin từ file..."):
                st.session_state.docx_text = extract_text_from_docx(uploaded_file)
                parsed_data = parse_info_from_text(st.session_state.docx_text)
                # Lưu dữ liệu đã trích xuất vào session_state
                for key, value in parsed_data.items():
                    st.session_state[key] = value
                st.session_state.data_extracted = True
                st.success("Trích xuất thông tin thành công! Vui lòng kiểm tra và hiệu chỉnh bên dưới.")
   
    # Sử dụng expander để nhóm các trường thông tin
    with st.expander("Vùng 1 - Thông tin khách hàng", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.ho_ten = st.text_input("Họ và tên", value=st.session_state.get('ho_ten', ''))
            st.session_state.cccd = st.text_input("CCCD/CMND", value=st.session_state.get('cccd', ''))
        with col2:
            st.session_state.sdt = st.text_input("Số điện thoại", value=st.session_state.get('sdt', ''))
            st.session_state.dia_chi = st.text_input("Địa chỉ", value=st.session_state.get('dia_chi', ''))
    with st.expander("Vùng 2 - Thông tin phương án vay", expanded=True):
        st.session_state.muc_dich_vay = st.text_area("Mục đích vay", value=st.session_state.get('muc_dich_vay', ''))
        col1, col2, col3 = st.columns(3)
        with col1:
            st.session_state.tong_nhu_cau_von = st.number_input(
                "Tổng nhu cầu vốn (VNĐ)",
                min_value=0,
                value=int(st.session_state.get('tong_nhu_cau_von', 0)),
                step=1000000,
                format="%d"
            )
            st.session_state.von_doi_ung = st.number_input(
                "Vốn đối ứng (VNĐ)",
                min_value=0,
                value=int(st.session_state.get('von_doi_ung', 0)),
                step=1000000,
                format="%d"
            )
        with col2:
            st.session_state.so_tien_vay = st.number_input(
                "Số tiền vay (VNĐ)",
                min_value=0,
                value=int(st.session_state.get('so_tien_vay', 0)),
                step=1000000,
                format="%d"
            )
            st.session_state.lai_suat = st.number_input(
                "Lãi suất (%/năm)",
                min_value=0.0,
                value=st.session_state.get('lai_suat', 5.0),
                step=0.1,
                format="%.1f"
            )
        with col3:
             st.session_state.thoi_gian_vay = st.number_input(
                "Thời gian vay (tháng)",
                min_value=1,
                value=int(st.session_state.get('thoi_gian_vay', 12)),
                step=1
            )
           
    with st.expander("Vùng 3 - Thông tin tài sản đảm bảo", expanded=True):
        st.session_state.tsdb_mo_ta = st.text_area("Mô tả tài sản", value=st.session_state.get('tsdb_mo_ta', ''))
        st.session_state.tsdb_gia_tri = st.number_input(
            "Giá trị định giá (VNĐ)",
            min_value=0,
            value=int(st.session_state.get('tsdb_gia_tri', 0)),
            step=1000000,
            format="%d"
        )
# Lấy dữ liệu từ session_state để tính toán
so_tien_vay = st.session_state.get('so_tien_vay', 0)
tong_nhu_cau_von = st.session_state.get('tong_nhu_cau_von', 0)
von_doi_ung = st.session_state.get('von_doi_ung', 0)
tsdb_gia_tri = st.session_state.get('tsdb_gia_tri', 0)
lai_suat = st.session_state.get('lai_suat', 0.0)
thoi_gian_vay = st.session_state.get('thoi_gian_vay', 0)
# --------------------------------------------------------------------------------------
# TAB 2: PHÂN TÍCH CHỈ SỐ & DÒNG TIỀN
# --------------------------------------------------------------------------------------
with tab2:
    st.header("Các chỉ số tài chính và Kế hoạch trả nợ")
    if so_tien_vay > 0:
        st.subheader("Các chỉ số tài chính quan trọng")
        col1, col2, col3 = st.columns(3)
       
        # Tính toán chỉ số
        ty_le_vay_tong_von = (so_tien_vay / tong_nhu_cau_von * 100) if tong_nhu_cau_von > 0 else 0
        ty_le_doi_ung = (von_doi_ung / tong_nhu_cau_von * 100) if tong_nhu_cau_von > 0 else 0
        ty_le_vay_tsdb = (so_tien_vay / tsdb_gia_tri * 100) if tsdb_gia_tri > 0 else 0
        # Lưu chỉ số để xuất báo cáo
        ratios_for_report = {
            "Tỷ lệ Vay/Tổng nhu cầu vốn": f"{ty_le_vay_tong_von:.2f}%",
            "Tỷ lệ Vốn đối ứng/Tổng nhu cầu vốn": f"{ty_le_doi_ung:.2f}%",
            "Tỷ lệ Vay/Giá trị TSĐB": f"{ty_le_vay_tsdb:.2f}%"
        }
        st.session_state.ratios = ratios_for_report
        col1.metric("Tỷ lệ Vay/Tổng nhu cầu vốn", f"{ty_le_vay_tong_von:.2f}%")
        col2.metric("Tỷ lệ Vốn đối ứng", f"{ty_le_doi_ung:.2f}%")
        col3.metric("Tỷ lệ Vay/Giá trị TSĐB", f"{ty_le_vay_tsdb:.2f}%", help="Tỷ lệ giữa số tiền vay và tổng giá trị tài sản đảm bảo.")
        st.markdown("---")
       
        st.subheader("Bảng kế hoạch trả nợ chi tiết")
        repayment_df = calculate_repayment_schedule(so_tien_vay, lai_suat, thoi_gian_vay)
       
        # Lưu bảng để có thể xuất Excel
        st.session_state.repayment_df = repayment_df
        if not repayment_df.empty:
            # Định dạng lại các cột tiền tệ để hiển thị
            df_display = repayment_df.copy()
            currency_cols = ["Dư nợ đầu kỳ", "Gốc trả trong kỳ", "Lãi trả trong kỳ", "Tổng gốc và lãi", "Dư nợ cuối kỳ"]
            for col in currency_cols:
                df_display[col] = df_display[col].apply(format_currency)
           
            st.dataframe(df_display, use_container_width=True, height=min(35 * (len(df_display) + 1), 600))
        else:
            st.warning("Vui lòng nhập đầy đủ thông tin khoản vay (Số tiền, Lãi suất, Thời gian) để xem kế hoạch trả nợ.")
    else:
        st.info("Chưa có dữ liệu để phân tích. Vui lòng nhập thông tin ở tab 'Nhập liệu'.")
# --------------------------------------------------------------------------------------
# TAB 3: BIỂU ĐỒ TRỰC QUAN
# --------------------------------------------------------------------------------------
with tab3:
    st.header("Trực quan hóa dữ liệu tài chính")
    if so_tien_vay > 0 and von_doi_ung > 0:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Cơ cấu Nguồn vốn")
            labels = ['Vốn vay', 'Vốn đối ứng']
            values = [so_tien_vay, von_doi_ung]
            fig_pie = go.Figure(data=[go.Pie(
                labels=labels,
                values=values,
                hole=.3,
                textinfo='percent+label',
                marker_colors=px.colors.sequential.Blues_r
            )])
            fig_pie.update_layout(
                title_text='Tỷ trọng Vốn vay và Vốn đối ứng',
                legend_title_text='Nguồn vốn',
                uniformtext_minsize=12,
                uniformtext_mode='hide'
            )
            st.plotly_chart(fig_pie, use_container_width=True)
        with col2:
            st.subheader("Biến động Dư nợ")
            if 'repayment_df' in st.session_state and not st.session_state.repayment_df.empty:
                df_chart = st.session_state.repayment_df
                # Thêm dòng dư nợ ban đầu tại kỳ 0
                initial_row = pd.DataFrame([{'Kỳ trả nợ': 0, 'Dư nợ cuối kỳ': so_tien_vay}])
                df_chart = pd.concat([initial_row, df_chart[['Kỳ trả nợ', 'Dư nợ cuối kỳ']]], ignore_index=True)
                fig_line = go.Figure()
                fig_line.add_trace(go.Scatter(
                    x=df_chart['Kỳ trả nợ'],
                    y=df_chart['Dư nợ cuối kỳ'],
                    mode='lines+markers',
                    name='Dư nợ',
                    fill='tozeroy' # Tô màu vùng dưới đường line
                ))
                fig_line.update_layout(
                    title_text='Dư nợ giảm dần qua các kỳ',
                    xaxis_title='Kỳ trả nợ (Tháng)',
                    yaxis_title='Dư nợ còn lại (VNĐ)',
                )
                st.plotly_chart(fig_line, use_container_width=True)
            else:
                 st.warning("Không có dữ liệu kế hoạch trả nợ để vẽ biểu đồ.")
    else:
        st.info("Chưa có dữ liệu để vẽ biểu đồ. Vui lòng nhập thông tin ở tab 'Nhập liệu'.")
# --------------------------------------------------------------------------------------
# TAB 4: PHÂN TÍCH BỞI AI
# --------------------------------------------------------------------------------------
with tab4:
    st.header("Phân tích Chuyên sâu với Trí tuệ Nhân tạo (Gemini)")
   
    if not api_key:
        st.warning("Vui lòng nhập Gemini API Key ở thanh bên để sử dụng tính năng này.")
    else:
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.5-flash') # Hoặc gemini-pro
        except Exception as e:
            st.error(f"Lỗi khởi tạo Gemini: {e}")
            model = None
        if st.button("🚀 Bắt đầu Phân tích", use_container_width=True, disabled=(not model or not st.session_state.data_extracted)):
           
            # Phân tích 1: Dựa trên file .docx gốc
            with st.spinner("AI đang phân tích nội dung file .docx..."):
                if st.session_state.docx_text:
                    prompt1 = f"""
                    Với vai trò là một chuyên gia thẩm định tín dụng ngân hàng, hãy phân tích toàn bộ nội dung của phương án kinh doanh dưới đây.
                    Cần tập trung vào các điểm sau:
                    1. Tóm tắt tổng quan về phương án kinh doanh.
                    2. Phân tích các điểm mạnh của phương án (ví dụ: kinh nghiệm, thị trường, sản phẩm).
                    3. Phân tích các điểm yếu hoặc các điểm cần làm rõ.
                    4. Nhận diện các rủi ro tiềm ẩn (thị trường, hoạt động, tài chính).
                    5. Đưa ra một kết luận sơ bộ về tính khả thi của phương án.
                    Nội dung phương án kinh doanh:
                    ---
                    {st.session_state.docx_text}
                    ---
                    """
                    try:
                        response1 = model.generate_content(prompt1)
                        st.session_state.ai_analysis_1 = response1.text
                    except Exception as e:
                        st.session_state.ai_analysis_1 = f"Lỗi khi gọi API Gemini: {e}"
                else:
                    st.session_state.ai_analysis_1 = "Không có nội dung file để phân tích."
           
            # Phân tích 2: Dựa trên dữ liệu đã hiệu chỉnh
            with st.spinner("AI đang phân tích các chỉ số và dữ liệu đã hiệu chỉnh..."):
                # Tổng hợp thông tin từ session_state thành một chuỗi
                adjusted_data_summary = f"""
                - Khách hàng: {st.session_state.ho_ten}, CCCD: {st.session_state.cccd}
                - Mục đích vay: {st.session_state.muc_dich_vay}
                - Tổng nhu cầu vốn: {format_currency(st.session_state.tong_nhu_cau_von)} VNĐ
                - Vốn đối ứng: {format_currency(st.session_state.von_doi_ung)} VNĐ
                - Số tiền vay: {format_currency(st.session_state.so_tien_vay)} VNĐ
                - Lãi suất: {st.session_state.lai_suat}%/năm
                - Thời gian vay: {st.session_state.thoi_gian_vay} tháng
                - Tổng giá trị TSĐB: {format_currency(st.session_state.tsdb_gia_tri)} VNĐ
                - Tỷ lệ Vay/TSĐB: {st.session_state.get('ratios', {}).get('Tỷ lệ Vay/Giá trị TSĐB', 'N/A')}
                - Tỷ lệ Vay/Tổng vốn: {st.session_state.get('ratios', {}).get('Tỷ lệ Vay/Tổng nhu cầu vốn', 'N/A')}
                """
               
                prompt2 = f"""
                Với vai trò là một chuyên gia thẩm định tín dụng, hãy phân tích sâu về các chỉ số tài chính của phương án vay vốn dựa trên các thông số đã được chuyên viên tín dụng hiệu chỉnh dưới đây.
                Hãy tập trung vào:
                1. Đánh giá tính hợp lý của số tiền vay so với nhu cầu vốn và vốn đối ứng.
                2. Phân tích khả năng trả nợ dựa trên số tiền vay, lãi suất và thời hạn.
                3. Đánh giá mức độ an toàn của khoản vay dựa trên tỷ lệ cho vay so với giá trị tài sản đảm bảo.
                4. Đưa ra các khuyến nghị (nếu có) để tăng tính khả thi cho phương án.
                Dữ liệu đã hiệu chỉnh:
                ---
                {adjusted_data_summary}
                ---
                """
                try:
                    response2 = model.generate_content(prompt2)
                    st.session_state.ai_analysis_2 = response2.text
                except Exception as e:
                     st.session_state.ai_analysis_2 = f"Lỗi khi gọi API Gemini: {e}"
        if 'ai_analysis_1' in st.session_state:
            with st.expander("1. Phân tích từ file .docx của khách hàng", expanded=True):
                st.info("Nguồn dữ liệu: Phân tích từ file .docx của khách hàng.")
                st.markdown(st.session_state.ai_analysis_1)
        if 'ai_analysis_2' in st.session_state:
            with st.expander("2. Phân tích từ các thông số và chỉ số đã tính toán", expanded=True):
                st.info("Nguồn dữ liệu: Phân tích từ các thông số và chỉ số đã tính toán trên ứng dụng.")
                st.markdown(st.session_state.ai_analysis_2)
# --------------------------------------------------------------------------------------
# TAB 5: CHATBOT HỖ TRỢ
# --------------------------------------------------------------------------------------
with tab5:
    st.header("Chatbot Hỗ trợ Thẩm định")
    if not api_key:
        st.warning("Vui lòng nhập Gemini API Key ở thanh bên để sử dụng tính năng này.")
    else:
        try:
            # Khởi tạo model cho chatbot
            model_chat = genai.GenerativeModel('gemini-2.5-flash')
            chat = model_chat.start_chat(history=[])
        except Exception as e:
            st.error(f"Lỗi khởi tạo Gemini Chat: {e}")
            chat = None
        # Hiển thị lịch sử chat
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        # Input từ người dùng
        if prompt := st.chat_input("Đặt câu hỏi về thẩm định, tài chính..."):
            if chat:
                # Thêm tin nhắn của người dùng vào lịch sử và hiển thị
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)
                # Gửi tin nhắn đến Gemini và nhận phản hồi
                with st.chat_message("assistant"):
                    with st.spinner("Bot đang suy nghĩ..."):
                        try:
                            # Bổ sung context từ file vào prompt
                            context_prompt = f"""
                            Dựa trên bối cảnh của phương án kinh doanh này (nếu có):
                            ---
                            {st.session_state.docx_text[:2000]}...
                            ---
                            Và dữ liệu tổng hợp:
                            ---
                            - Khách hàng: {st.session_state.get('ho_ten', 'N/A')}
                            - Số tiền vay: {format_currency(st.session_state.get('so_tien_vay', 0))} VNĐ
                            - Mục đích: {st.session_state.get('muc_dich_vay', 'N/A')}
                            ---
                            Hãy trả lời câu hỏi sau: "{prompt}"
                            """
                            response = chat.send_message(context_prompt)
                            response_text = response.text
                            st.markdown(response_text)
                            # Thêm phản hồi của bot vào lịch sử
                            st.session_state.chat_history.append({"role": "assistant", "content": response_text})
                        except Exception as e:
                            error_message = f"Xin lỗi, đã có lỗi xảy ra: {e}"
                            st.error(error_message)
                            st.session_state.chat_history.append({"role": "assistant", "content": error_message})
            else:
                st.error("Không thể khởi tạo chatbot. Vui lòng kiểm tra API Key.")
        if st.session_state.chat_history:
            if st.button("🗑️ Xóa lịch sử trò chuyện"):
                st.session_state.chat_history = []
                st.rerun()
# ======================================================================================
# LOGIC XỬ LÝ NÚT EXPORT (đặt ở cuối để truy cập được mọi state)
# ======================================================================================
if execute_export:
    if export_option == "Xuất Kế hoạch trả nợ (Excel)":
        if 'repayment_df' in st.session_state and not st.session_state.repayment_df.empty:
            excel_data = generate_excel_download(st.session_state.repayment_df)
            st.sidebar.download_button(
                label="📥 Tải xuống file Excel",
                data=excel_data,
                file_name=f"KeHoachTraNo_{st.session_state.get('ho_ten', 'KH')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            st.sidebar.success("Đã tạo file Excel!")
        else:
            st.sidebar.error("Không có dữ liệu kế hoạch trả nợ để xuất.")
    elif export_option == "Xuất Báo cáo Thẩm định":
        if st.session_state.data_extracted:
            with st.spinner("Đang tạo báo cáo..."):
                customer_info = {k: st.session_state.get(k) for k in ['ho_ten', 'cccd', 'sdt', 'dia_chi']}
                loan_info = {k: st.session_state.get(k) for k in ['muc_dich_vay', 'tong_nhu_cau_von', 'von_doi_ung', 'so_tien_vay', 'lai_suat', 'thoi_gian_vay']}
                collateral_info = {k: st.session_state.get(k) for k in ['tsdb_mo_ta', 'tsdb_gia_tri']}
               
                report_data = generate_report_docx(
                    customer_info,
                    loan_info,
                    collateral_info,
                    st.session_state.get('ratios', {}),
                    st.session_state.get('ai_analysis_1', 'Chưa thực hiện phân tích.'),
                    st.session_state.get('ai_analysis_2', 'Chưa thực hiện phân tích.')
                )
                st.sidebar.download_button(
                    label="📥 Tải xuống Báo cáo (.docx)",
                    data=report_data,
                    file_name=f"BaoCaoThamDinh_{st.session_state.get('ho_ten', 'KH')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.sidebar.success("Đã tạo file báo cáo!")
        else:
            st.sidebar.error("Chưa có dữ liệu để tạo báo cáo.")
